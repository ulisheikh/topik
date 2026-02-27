import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from config import (
    WORD_COL_NUMBER_WIDTH,
    WORD_COL_QUESTION_WIDTH,
    WORD_COL_ANSWER_WIDTH,
    WORD_ROW_HEIGHT
)

def set_cell_border(cell, **kwargs):
    """Kataklar uchun qalin qora chiziq"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        if kwargs.get(side):
            edge = OxmlElement(f'w:{side}')
            edge.set(qn('w:val'), 'single')
            edge.set(qn('w:sz'), '8')
            edge.set(qn('w:space'), '0')
            edge.set(qn('w:color'), '000000')
            tcBorders.append(edge)
    tcPr.append(tcBorders)

def create_exam_word(words, location=None, mode="kr_to_uz", filename_prefix="exam"):
    """BARCHA so'zlarni bir faylda chiqarish - 20 tadan listlar bo'lib"""
    doc = Document()

    # A5 format
    section = doc.sections[0]
    section.page_height = Cm(21.0)
    section.page_width = Cm(14.8)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)

    # So'zlarni 20 tadan guruhlar bo'lib ajratish
    groups = split_words_into_groups(words)
    
    # Har bir guruh uchun alohida jadval yaratish
    for group_num, group_words in enumerate(groups, 1):
        # Agar birinchi guruh bo'lmasa, yangi sahifa qo'shish
        if group_num > 1:
            doc.add_page_break()
        
        # 1. ADRESS (📍 Mavzu)
        if location:
            addr_para = doc.add_paragraph()
            addr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            addr_run = addr_para.add_run(f"📍 {location} - List {group_num}/{len(groups)}")
            addr_run.font.size = Pt(11)
            addr_run.font.bold = True
            addr_para.paragraph_format.space_after = Pt(10)

        # 2. JADVAL YARATISH
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # USTUN O'LCHAMLARI (config dan)
        table.columns[0].width = Cm(WORD_COL_NUMBER_WIDTH)
        table.columns[1].width = Cm(WORD_COL_QUESTION_WIDTH)
        table.columns[2].width = Cm(WORD_COL_ANSWER_WIDTH)

        # 3. UNIVERSAL SARLAVHALAR
        header_cells = table.rows[0].cells
        header_names = ["번호", "질문", "답안"] 
        
        for i, name in enumerate(header_names):
            cell = header_cells[i]
            cell.text = name
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
            
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.bold = True
            run.font.size = Pt(11)

        # 4. SO'ZLARNI TO'LDIRISH
        for idx, (korean, uzbek) in enumerate(group_words, 1):
            row = table.add_row()
            row.height = Cm(WORD_ROW_HEIGHT)

            # № (Bold)
            row.cells[0].text = str(idx)
            
            # Savol ustuni (Mode ga qarab Koreyscha yoki O'zbekcha)
            if mode == "kr_to_uz":
                row.cells[1].text = korean
            else:
                row.cells[1].text = uzbek
                
            row.cells[2].text = ""  # Javob yozish uchun bo'sh joy

            # Formatlash
            for i, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_border(cell, top=True, bottom=True, left=True, right=True)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.size = Pt(10)
                    if i == 0:  # Raqamni bold qilish
                        p.runs[0].font.bold = True

    # Saqlash - ANIQ NOM BILAN
    temp_dir = "temp_exams"
    os.makedirs(temp_dir, exist_ok=True)
    filename = f"{filename_prefix}.docx"
    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)
    
    return filepath

def split_words_into_groups(words, words_per_file=20):
    groups = []
    for i in range(0, len(words), words_per_file):
        groups.append(words[i:i + words_per_file])
    return groups


def create_exam_word_bilingual(words, location=None, filename_prefix="bilingual"):
    """Har ikki tilda (한국어 | O'zbek) jadval - savol/javob yo'q, tarjima bilan birga"""
    doc = Document()

    # A5 format
    section = doc.sections[0]
    section.page_height = Cm(21.0)
    section.page_width = Cm(14.8)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(0.8)
    section.right_margin = Cm(0.8)

    groups = split_words_into_groups(words)

    for group_num, group_words in enumerate(groups, 1):
        if group_num > 1:
            doc.add_page_break()

        # Sarlavha
        if location:
            addr_para = doc.add_paragraph()
            addr_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            addr_run = addr_para.add_run(f"📍 {location} - List {group_num}/{len(groups)}")
            addr_run.font.size = Pt(11)
            addr_run.font.bold = True
            addr_para.paragraph_format.space_after = Pt(10)

        # Jadval: №  |  한국어  |  O'zbekcha
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Config dan o'lchamlar
        table.columns[0].width = Cm(WORD_COL_NUMBER_WIDTH)
        table.columns[1].width = Cm(WORD_COL_QUESTION_WIDTH)
        table.columns[2].width = Cm(WORD_COL_ANSWER_WIDTH)

        # Sarlavha qatori
        header_cells = table.rows[0].cells
        header_names = ["번호", "한국어", "O'zbekcha"]

        for i, name in enumerate(header_names):
            cell = header_cells[i]
            cell.text = name
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, top=True, bottom=True, left=True, right=True)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.font.bold = True
            run.font.size = Pt(11)

        # So'zlarni to'ldirish
        for idx, (korean, uzbek) in enumerate(group_words, 1):
            # * belgisini olib tashlash
            korean_clean = korean.lstrip('*') if isinstance(korean, str) else korean
            uzbek_clean = uzbek.lstrip('*') if isinstance(uzbek, str) else uzbek

            row = table.add_row()
            row.height = Cm(WORD_ROW_HEIGHT)

            row.cells[0].text = str(idx)
            row.cells[1].text = korean_clean
            row.cells[2].text = uzbek_clean

            for i, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_border(cell, top=True, bottom=True, left=True, right=True)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.size = Pt(10)
                    if i == 0:
                        p.runs[0].font.bold = True

    # Saqlash - ANIQ NOM BILAN
    temp_dir = "temp_exams"
    os.makedirs(temp_dir, exist_ok=True)
    filename = f"{filename_prefix}.docx"
    filepath = os.path.join(temp_dir, filename)
    doc.save(filepath)

    return filepath